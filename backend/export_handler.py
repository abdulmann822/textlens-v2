import os
from pathlib import Path


class ExportHandler:
    """Exports extracted text to TXT, PDF, or DOCX."""

    def export(self, text: str, filename: str, fmt: str, output_dir: str) -> str:
        # Sanitize filename
        safe_name = "".join(c for c in filename if c.isalnum() or c in "._- ").strip()
        if not safe_name:
            safe_name = "extracted_text"
        base = Path(safe_name).stem
        output_path = os.path.join(output_dir, f"{base}.{fmt}")

        if fmt == "txt":
            self._export_txt(text, output_path)
        elif fmt == "pdf":
            self._export_pdf(text, output_path, base)
        elif fmt == "docx":
            self._export_docx(text, output_path, base)
        else:
            raise ValueError(f"Unsupported format: {fmt}")

        return output_path

    def _export_txt(self, text: str, path: str):
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)

    def _export_pdf(self, text: str, path: str, title: str):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib import colors

        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=25 * mm, bottomMargin=25 * mm,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title2", parent=styles["Title"],
            fontSize=16, spaceAfter=12,
            textColor=colors.HexColor("#1e293b"),
        )
        body_style = ParagraphStyle(
            "Body2", parent=styles["Normal"],
            fontSize=11, leading=16,
            textColor=colors.HexColor("#374151"),
            alignment=TA_LEFT,
        )

        story = [Paragraph(title, title_style), Spacer(1, 6 * mm)]
        for para in text.split("\n\n"):
            para = para.strip()
            if not para:
                story.append(Spacer(1, 4 * mm))
                continue
            safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe = safe.replace("\n", "<br/>")
            story.append(Paragraph(safe, body_style))
            story.append(Spacer(1, 3 * mm))

        doc.build(story)

    def _export_docx(self, text: str, path: str, title: str):
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if heading.runs:
            heading.runs[0].font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

        doc.add_paragraph()

        for para in text.split("\n\n"):
            para = para.strip()
            if not para:
                doc.add_paragraph()
                continue
            if para.startswith("--- Page") and para.endswith("---"):
                h = doc.add_heading(para.strip("- ").strip(), level=2)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(para)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        doc.save(path)
